from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


WORK = Path(__file__).resolve().parents[1]
DOCX_PATH = WORK / "output" / "parabolic-motion-radar-effects-conference-paper.docx"


def main() -> None:
    doc = Document(DOCX_PATH)
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    headings = [p.text for p in paragraphs if p.style.name.startswith("Heading")]
    equations = [p.text for p in paragraphs if p.style.name == "Equation Block"]
    captions = [p.text for p in paragraphs if p.style.name == "Caption Custom"]
    tables = doc.tables

    with ZipFile(DOCX_PATH, "r") as zf:
        names = zf.namelist()
        media = [name for name in names if name.startswith("word/media/")]
        empty_media = [name for name in media if zf.getinfo(name).file_size == 0]
        document_xml = zf.read("word/document.xml")

    checks = {
        "paragraph_count": len(paragraphs),
        "heading_count": len(headings),
        "equation_count": len(equations),
        "caption_count": len(captions),
        "table_count": len(tables),
        "media_count": len(media),
        "empty_media_count": len(empty_media),
        "has_document_xml": bool(document_xml),
        "has_title": any("A Simplified Parabolic-Motion Model" in p.text for p in paragraphs),
        "has_abstract": "Abstract" in headings,
        "has_conclusion": "11. Conclusion" in headings,
    }

    for key, value in checks.items():
        print(f"{key}: {value}")

    for idx, table in enumerate(tables, start=1):
        print(f"table_{idx}_rows: {len(table.rows)}")
        print(f"table_{idx}_cols: {len(table.columns)}")

    if empty_media:
        raise SystemExit(f"Empty media files found: {empty_media}")
    if checks["media_count"] < 2:
        raise SystemExit("Expected at least two embedded media files.")
    if checks["table_count"] < 3:
        raise SystemExit("Expected at least three tables.")
    if checks["equation_count"] < 10:
        raise SystemExit("Expected at least ten equation blocks.")
    if not checks["has_title"] or not checks["has_abstract"] or not checks["has_conclusion"]:
        raise SystemExit("Missing required paper sections.")


if __name__ == "__main__":
    main()
