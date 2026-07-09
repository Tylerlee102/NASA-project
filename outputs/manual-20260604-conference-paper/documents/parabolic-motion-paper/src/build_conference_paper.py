from __future__ import annotations

import math
from pathlib import Path

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[5]
WORK = Path(__file__).resolve().parents[1]
FIG_DIR = WORK / "figures"
OUT_DIR = WORK / "output"
DOCX_PATH = OUT_DIR / "parabolic-motion-radar-effects-conference-paper.docx"


TOKENS = {
    "font": "Calibri",
    "mono_font": "Consolas",
    "body_size": 11,
    "title_size": 21,
    "h1_size": 16,
    "h2_size": 13,
    "h3_size": 12,
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "ink": "0B2545",
    "muted": "667085",
    "light_fill": "F2F4F7",
    "callout_fill": "F4F6F9",
    "border": "D0D5DD",
}


def hex_color(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, col_widths_dxa: list[int], indent_dxa: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(col_widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in col_widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def model_values() -> dict[str, float]:
    z0_km = 400.0
    y_km = 25.0
    x_edge_km = 60.0
    rise_km = 4.0
    speed_km_s = 4.0
    n_ice = 1.78
    c = 299_792_458.0
    vhf_lambda_m = 5.0
    hf_lambda_m = 33.3
    baseline_m = 5.0

    r_km = math.sqrt(z0_km**2 + y_km**2)
    extra_range_m = (r_km - z0_km) * 1000.0
    apparent_depth_m = extra_range_m / n_ice
    extra_delay_us = 2.0 * extra_range_m / c * 1.0e6
    look_angle_deg = math.degrees(math.atan2(y_km, z0_km))
    phase_rad = 2.0 * math.pi / vhf_lambda_m * baseline_m * math.sin(math.radians(look_angle_deg))
    phase_deg = math.degrees(phase_rad)
    power_db = 10.0 * math.log10((z0_km / r_km) ** 4)

    x = np.linspace(-60.0, 60.0, 241)
    a = rise_km / x_edge_km**2
    z = z0_km + a * x**2
    dz_dt = 2.0 * a * x * speed_km_s
    r = np.sqrt(x**2 + y_km**2 + z**2)
    range_rate_m_s = 1000.0 * (x * speed_km_s + z * dz_dt) / r
    vhf_doppler = -2.0 * range_rate_m_s / vhf_lambda_m
    hf_doppler = -2.0 * range_rate_m_s / hf_lambda_m

    return {
        "z0_km": z0_km,
        "y_km": y_km,
        "x_edge_km": x_edge_km,
        "rise_km": rise_km,
        "speed_km_s": speed_km_s,
        "n_ice": n_ice,
        "c": c,
        "vhf_lambda_m": vhf_lambda_m,
        "hf_lambda_m": hf_lambda_m,
        "baseline_m": baseline_m,
        "extra_range_m": extra_range_m,
        "apparent_depth_m": apparent_depth_m,
        "extra_delay_us": extra_delay_us,
        "look_angle_deg": look_angle_deg,
        "phase_deg": phase_deg,
        "power_db": power_db,
        "max_vhf_doppler_hz": float(np.max(np.abs(vhf_doppler))),
        "max_hf_doppler_hz": float(np.max(np.abs(hf_doppler))),
        "min_prf_hz": float(2.0 * np.max(np.abs(vhf_doppler))),
    }


def chart_font(size: int, bold: bool = False):
    candidates = ["arialbd.ttf" if bold else "arial.ttf", "calibrib.ttf" if bold else "calibri.ttf"]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_center(draw: ImageDraw.ImageDraw, xy: tuple[float, float], text: str, font, fill: str) -> None:
    box = draw.textbbox((0, 0), text, font=font)
    draw.text((xy[0] - (box[2] - box[0]) / 2, xy[1] - (box[3] - box[1]) / 2), text, font=font, fill=fill)


def data_to_px(
    x: float,
    y: float,
    plot_box: tuple[int, int, int, int],
    xlim: tuple[float, float],
    ylim: tuple[float, float],
) -> tuple[float, float]:
    left, top, right, bottom = plot_box
    px = left + (x - xlim[0]) / (xlim[1] - xlim[0]) * (right - left)
    py = bottom - (y - ylim[0]) / (ylim[1] - ylim[0]) * (bottom - top)
    return px, py


def draw_polyline(
    draw: ImageDraw.ImageDraw,
    xs: np.ndarray,
    ys: np.ndarray,
    plot_box: tuple[int, int, int, int],
    xlim: tuple[float, float],
    ylim: tuple[float, float],
    fill: str,
    width: int = 4,
) -> None:
    points = [data_to_px(float(x), float(y), plot_box, xlim, ylim) for x, y in zip(xs, ys)]
    draw.line(points, fill=fill, width=width, joint="curve")


def draw_axes(
    draw: ImageDraw.ImageDraw,
    plot_box: tuple[int, int, int, int],
    xlim: tuple[float, float],
    ylim: tuple[float, float],
    title: str,
    xlabel: str,
    ylabel: str,
    xticks: list[float],
    yticks: list[float],
) -> None:
    left, top, right, bottom = plot_box
    axis = "#344054"
    grid = "#E4E7EC"
    label_font = chart_font(23)
    tick_font = chart_font(20)
    title_font = chart_font(25, bold=True)

    draw.rectangle(plot_box, fill="white", outline="#D0D5DD", width=2)
    for xt in xticks:
        px, _ = data_to_px(xt, ylim[0], plot_box, xlim, ylim)
        draw.line((px, top, px, bottom), fill=grid, width=1)
        draw.line((px, bottom, px, bottom + 7), fill=axis, width=2)
        text_center(draw, (px, bottom + 25), f"{xt:g}", tick_font, "#475467")
    for yt in yticks:
        _, py = data_to_px(xlim[0], yt, plot_box, xlim, ylim)
        draw.line((left, py, right, py), fill=grid, width=1)
        draw.line((left - 7, py, left, py), fill=axis, width=2)
        draw.text((left - 72, py - 10), f"{yt:g}", font=tick_font, fill="#475467")

    draw.line((left, bottom, right, bottom), fill=axis, width=2)
    draw.line((left, top, left, bottom), fill=axis, width=2)
    text_center(draw, ((left + right) / 2, top - 31), title, title_font, "#0B2545")
    text_center(draw, ((left + right) / 2, bottom + 58), xlabel, label_font, "#344054")
    draw.text((left, top - 22), ylabel, font=label_font, fill="#344054")


def make_figures(values: dict[str, float]) -> tuple[Path, Path]:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    x = np.linspace(-60.0, 60.0, 241)
    z = values["z0_km"] + values["rise_km"] * (x / values["x_edge_km"]) ** 2
    y = values["y_km"]
    n_ice = values["n_ice"]
    c = values["c"]
    baseline_m = values["baseline_m"]
    lambda_v = values["vhf_lambda_m"]
    lambda_h = values["hf_lambda_m"]
    speed = values["speed_km_s"]

    geometry_path = FIG_DIR / "figure_1_parabolic_motion_geometry.png"
    img = Image.new("RGB", (1800, 760), "white")
    draw = ImageDraw.Draw(img)
    title_font = chart_font(32, bold=True)
    label_font = chart_font(23)
    note_font = chart_font(22)
    draw.text((70, 35), "Simplified parabolic-motion geometry", font=title_font, fill="#0B2545")

    box1 = (125, 145, 825, 590)
    box2 = (1030, 145, 1660, 590)
    draw_axes(
        draw,
        box1,
        (-60, 60),
        (399.5, 404.7),
        "Parabolic altitude model",
        "along-track x (km)",
        "altitude z (km)",
        [-60, -30, 0, 30, 60],
        [400, 401, 402, 403, 404],
    )
    draw_polyline(draw, x, z, box1, (-60, 60), (399.5, 404.7), "#2E74B5", width=6)
    ca = data_to_px(0, values["z0_km"], box1, (-60, 60), (399.5, 404.7))
    draw.ellipse((ca[0] - 10, ca[1] - 10, ca[0] + 10, ca[1] + 10), fill="#0B2545")
    draw.line((ca[0] + 10, ca[1] - 4, ca[0] + 145, ca[1] - 110), fill="#475467", width=3)
    draw.text((ca[0] + 150, ca[1] - 130), "closest approach\nz0 = 400 km", font=note_font, fill="#475467")

    draw_axes(
        draw,
        box2,
        (-60, 60),
        (-8, 35),
        "Plan-view side offset",
        "along-track x (km)",
        "cross-track y (km)",
        [-60, -30, 0, 30, 60],
        [0, 10, 20, 30],
    )
    draw_polyline(draw, x, np.zeros_like(x), box2, (-60, 60), (-8, 35), "#98A2B3", width=5)
    npt = data_to_px(0, 0, box2, (-60, 60), (-8, 35))
    fpt = data_to_px(0, y, box2, (-60, 60), (-8, 35))
    draw.line((npt[0], npt[1], fpt[0], fpt[1]), fill="#B54708", width=5)
    draw.ellipse((npt[0] - 8, npt[1] - 8, npt[0] + 8, npt[1] + 8), fill="#0B2545")
    draw.ellipse((fpt[0] - 12, fpt[1] - 12, fpt[0] + 12, fpt[1] + 12), fill="#B54708")
    draw.line((fpt[0] + 10, fpt[1] + 4, fpt[0] + 140, fpt[1] + 72), fill="#B54708", width=3)
    draw.text((fpt[0] + 150, fpt[1] + 48), "example feature\nside offset y = 25 km", font=note_font, fill="#B54708")
    draw.text((70, 700), "Note: the altitude axis is compressed so the local curvature can be seen.", font=label_font, fill="#667085")
    img.save(geometry_path)

    r_km = np.sqrt(x**2 + y**2 + z**2)
    apparent_depth = ((r_km - z) * 1000.0) / n_ice
    extra_delay = 2.0 * ((r_km - z) * 1000.0) / c * 1.0e6
    horizontal_km = np.sqrt(x**2 + y**2)
    theta = np.arctan2(horizontal_km, z)
    phase_deg = np.degrees(2.0 * np.pi / lambda_v * baseline_m * np.sin(theta))
    a = values["rise_km"] / values["x_edge_km"] ** 2
    dz_dt = 2.0 * a * x * speed
    range_rate_m_s = 1000.0 * (x * speed + z * dz_dt) / r_km
    vhf_d = -2.0 * range_rate_m_s / lambda_v
    hf_d = -2.0 * range_rate_m_s / lambda_h

    results_path = FIG_DIR / "figure_2_compact_observable_response.png"
    img = Image.new("RGB", (1500, 1850), "white")
    draw = ImageDraw.Draw(img)
    draw.text((85, 42), "Compact observable response across the modeled pass", font=chart_font(33, bold=True), fill="#0B2545")
    boxes = [(165, 175, 1390, 610), (165, 750, 1390, 1185), (165, 1325, 1390, 1760)]
    draw_axes(
        draw,
        boxes[0],
        (-60, 60),
        (0, 3200),
        "Extra slant path plotted as apparent depth",
        "",
        "apparent depth (m)",
        [-60, -30, 0, 30, 60],
        [0, 800, 1600, 2400, 3200],
    )
    draw_polyline(draw, x, apparent_depth, boxes[0], (-60, 60), (0, 3200), "#2E74B5", width=6)
    p0 = data_to_px(0, apparent_depth[len(apparent_depth) // 2], boxes[0], (-60, 60), (0, 3200))
    draw.ellipse((p0[0] - 9, p0[1] - 9, p0[0] + 9, p0[1] + 9), fill="#0B2545")
    draw.text((p0[0] + 22, p0[1] - 18), f"{values['apparent_depth_m']:.0f} m at closest approach", font=chart_font(21), fill="#0B2545")

    draw_axes(
        draw,
        boxes[1],
        (-60, 60),
        (0, 65),
        "Delay and VHF phase inherit the same range geometry",
        "",
        "us / degrees",
        [-60, -30, 0, 30, 60],
        [0, 15, 30, 45, 60],
    )
    draw_polyline(draw, x, extra_delay, boxes[1], (-60, 60), (0, 65), "#175CD3", width=5)
    draw_polyline(draw, x, phase_deg, boxes[1], (-60, 60), (0, 65), "#B54708", width=5)
    draw.rectangle((1010, boxes[1][1] + 18, 1345, boxes[1][1] + 86), fill="white", outline="#E4E7EC")
    draw.line((1030, boxes[1][1] + 41, 1090, boxes[1][1] + 41), fill="#175CD3", width=6)
    draw.text((1105, boxes[1][1] + 28), "extra delay (us)", font=chart_font(21), fill="#175CD3")
    draw.line((1030, boxes[1][1] + 71, 1090, boxes[1][1] + 71), fill="#B54708", width=6)
    draw.text((1105, boxes[1][1] + 58), "VHF phase (deg)", font=chart_font(21), fill="#B54708")

    draw_axes(
        draw,
        boxes[2],
        (-60, 60),
        (-500, 500),
        "Range rate creates wavelength-dependent Doppler",
        "along-track x (km)",
        "Doppler (Hz)",
        [-60, -30, 0, 30, 60],
        [-500, -250, 0, 250, 500],
    )
    draw_polyline(draw, x, vhf_d, boxes[2], (-60, 60), (-500, 500), "#2E74B5", width=5)
    draw_polyline(draw, x, hf_d, boxes[2], (-60, 60), (-500, 500), "#667085", width=5)
    draw.rectangle((1010, boxes[2][1] + 18, 1330, boxes[2][1] + 62), fill="white", outline="#E4E7EC")
    draw.line((1030, boxes[2][1] + 41, 1090, boxes[2][1] + 41), fill="#2E74B5", width=6)
    draw.text((1105, boxes[2][1] + 28), "VHF", font=chart_font(21), fill="#2E74B5")
    draw.line((1180, boxes[2][1] + 41, 1240, boxes[2][1] + 41), fill="#667085", width=6)
    draw.text((1255, boxes[2][1] + 28), "HF", font=chart_font(21), fill="#667085")
    img.save(results_path)
    return geometry_path, results_path


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION_START.NEW_PAGE

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = TOKENS["font"]
    normal.font.size = Pt(TOKENS["body_size"])
    normal.font.color.rgb = hex_color(TOKENS["ink"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", TOKENS["h1_size"], TOKENS["blue"], 16, 8),
        ("Heading 2", TOKENS["h2_size"], TOKENS["blue"], 12, 6),
        ("Heading 3", TOKENS["h3_size"], TOKENS["dark_blue"], 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = TOKENS["font"]
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = hex_color(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title_style = styles.add_style("Paper Title", 1)
    title_style.font.name = TOKENS["font"]
    title_style.font.size = Pt(TOKENS["title_size"])
    title_style.font.bold = True
    title_style.font.color.rgb = hex_color(TOKENS["ink"])
    title_style.paragraph_format.space_after = Pt(3)
    title_style.paragraph_format.line_spacing = 1.05

    subtitle_style = styles.add_style("Paper Subtitle", 1)
    subtitle_style.font.name = TOKENS["font"]
    subtitle_style.font.size = Pt(11)
    subtitle_style.font.italic = True
    subtitle_style.font.color.rgb = hex_color(TOKENS["muted"])
    subtitle_style.paragraph_format.space_after = Pt(10)

    caption = styles.add_style("Caption Custom", 1)
    caption.font.name = TOKENS["font"]
    caption.font.size = Pt(9)
    caption.font.color.rgb = hex_color(TOKENS["muted"])
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.line_spacing = 1.05

    equation = styles.add_style("Equation Block", 1)
    equation.font.name = TOKENS["mono_font"]
    equation.font.size = Pt(9)
    equation.font.color.rgb = hex_color(TOKENS["ink"])
    equation.paragraph_format.left_indent = Inches(0.18)
    equation.paragraph_format.right_indent = Inches(0.1)
    equation.paragraph_format.space_before = Pt(2)
    equation.paragraph_format.space_after = Pt(4)
    equation.paragraph_format.line_spacing = 1.0

    small = styles.add_style("Small Body", 1)
    small.font.name = TOKENS["font"]
    small.font.size = Pt(9.3)
    small.font.color.rgb = hex_color(TOKENS["ink"])
    small.paragraph_format.space_after = Pt(4)
    small.paragraph_format.line_spacing = 1.05

    header = section.header.paragraphs[0]
    header.text = "Parabolic-motion radar-effects model"
    header.style = styles["Small Body"]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer = section.footer.paragraphs[0]
    footer.style = styles["Small Body"]
    add_page_number(footer)


def add_title_block(doc: Document) -> None:
    title = doc.add_paragraph(style="Paper Title")
    title.add_run("A Simplified Parabolic-Motion Model for Interpreting Radar Sounding Observables During a Flyby")
    subtitle = doc.add_paragraph(style="Paper Subtitle")
    subtitle.add_run("Example conference research paper draft prepared from the local NASA/REASON radar-analysis project")
    meta = doc.add_paragraph(style="Small Body")
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.add_run("Author: ").bold = True
    meta.add_run("Student Name (placeholder), Course / Institution (placeholder)\n")
    meta.add_run("Date: ").bold = True
    meta.add_run("June 4, 2026\n")
    meta.add_run("Scope note: ").bold = True
    meta.add_run(
        "This draft studies geometry-driven radar effects from modeled parabolic motion. "
        "It is not a claim that a specific subsurface water body or off-nadir reflector has been proven."
    )


def add_callout(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Small Body"]
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, TOKENS["callout_fill"])
    run = p.add_run(label + ": ")
    run.bold = True
    run.font.color.rgb = hex_color(TOKENS["dark_blue"])
    p.add_run(text)


def add_equation(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Equation Block")
    shade_paragraph(p, "F8FAFC")
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths)

    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        cell = hdr[idx]
        set_cell_shading(cell, TOKENS["light_fill"])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.3)
        run.font.color.rgb = hex_color(TOKENS["ink"])

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx != 1 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.size = Pt(9)
            run.font.color.rgb = hex_color(TOKENS["ink"])
    set_table_geometry(table, widths)
    doc.add_paragraph("", style="Small Body")


def add_figure(doc: Document, path: Path, caption: str, width_inches: float = 6.25) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph(style="Caption Custom")
    cap.add_run(caption)


def write_paper(doc: Document, values: dict[str, float], geometry: Path, results: Path) -> None:
    add_title_block(doc)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "Radar sounding measurements are commonly interpreted through time delay, range, phase, and Doppler. "
        "During a flyby, however, the observing platform does not remain at a constant position: altitude, "
        "slant range, and radial velocity change throughout the pass. This example paper develops a simplified "
        "parabolic-motion model to show how those changing geometry terms propagate into radar observables. "
        "The model uses a local along-track coordinate x, a fixed cross-track offset y, and a parabolic altitude "
        "profile z(x). From this geometry, the paper derives equations for slant range, apparent depth, two-way "
        "echo delay, interferometric phase, relative power, Doppler shift, and pulse-repetition-frequency (PRF) "
        "requirements. A representative case uses z0 = 400 km, y = 25 km, a 4 km altitude rise across +/-60 km, "
        "a platform speed of 4 km/s, an ice refractive index of 1.78, and REASON-like HF/VHF wavelengths. "
        f"In this case the 25 km side offset produces about {values['apparent_depth_m']:.1f} m of apparent depth, "
        f"{values['extra_delay_us']:.2f} us of extra two-way delay, and about {values['phase_deg']:.1f} degrees "
        "of VHF interferometric phase at closest approach. The modeled VHF Doppler magnitude reaches roughly "
        f"{values['max_vhf_doppler_hz']:.0f} Hz, implying a simple Nyquist PRF floor near "
        f"{values['min_prf_hz']:.0f} Hz for that geometry. The main expected outcome is that radar traces, "
        "phase, and Doppler change smoothly and predictably with the parabolic path; the model is a geometry "
        "interpretation tool, not a proof of a particular subsurface composition."
    )

    doc.add_paragraph(
        "Keywords: radar sounding; parabolic motion; flyby geometry; apparent depth; Doppler; PRF; interferometry; REASON."
    )
    add_callout(
        doc,
        "Research viewpoint",
        "The problem is to predict how radar measurements change when the observing object moves along a parabolic path. "
        "The useful question is therefore not just where an echo appears, but why the apparent depth, delay, phase, "
        "and Doppler all change together."
    )

    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Flyby radar sounding is a geometry-sensitive measurement. A transmitted pulse travels from the spacecraft "
        "to a target and back, and the receiver stores the return as a time delay that is later converted into "
        "range or apparent depth. If the spacecraft altitude and range are changing during the collection, then "
        "the measured delay also changes even when the target itself is fixed. This is the basic reason that a "
        "moving radar platform can create curved, shifted, or time-varying signatures in a radargram."
    )
    doc.add_paragraph(
        "The local project that motivated this paper began with slide-based calculations for a REASON-like radar "
        "case. Early versions treated a 25 km side offset and an approximately 440 m apparent-depth signature as "
        "if the goal were to decide whether an echo represented water or clutter. The corrected problem is broader "
        "and cleaner: determine the effects of parabolic motion on the radar observables. Once the motion model is "
        "defined, all of the radar quantities can be derived from the same range equation."
    )
    doc.add_paragraph(
        "This paper presents that corrected viewpoint as an example conference-style research draft. The emphasis "
        "is on explaining the chain of reasoning: define the path, compute slant range, convert range differences "
        "to delay and apparent depth, then examine phase, Doppler, and PRF implications. The result is a compact "
        "framework that can be used in a presentation or extended later with a real spacecraft trajectory."
    )

    doc.add_heading("2. Background and Motivation", level=1)
    doc.add_paragraph(
        "The REASON instrument on Europa Clipper provides a useful motivation because it is a dual-frequency radar "
        "sounder designed for active radar measurements over a wide flyby-altitude range. The local extraction from "
        "the REASON paper lists HF operation near 9 MHz with an approximately 33.3 m wavelength and VHF operation "
        "near 60 MHz with an approximately 5 m wavelength. It also lists operational altitudes spanning roughly "
        "25 km to 1000 km and a PRF range of about 50 Hz to 3000 Hz. Those values make the paper relevant for "
        "thinking about delay, resolution, Doppler, and PRF, but they do not provide the exact parabolic trajectory "
        "used in this simplified model."
    )
    doc.add_paragraph(
        "The distinction matters. The parabolic path in this paper is an idealized local approximation, not a "
        "mission-navigation product. It lets us ask a controlled question: if the platform altitude follows a "
        "parabola over a short along-track window, what radar effects should appear? A real mission analysis would "
        "replace this parabola with SPICE trajectory kernels, spacecraft attitude, antenna pointing, surface shape, "
        "and a scattering model."
    )

    doc.add_heading("3. Problem Statement", level=1)
    doc.add_paragraph(
        "The problem being solved is to predict the radar consequences of parabolic motion. The object or spacecraft "
        "is assumed to move along track while its altitude changes as a parabola. The model then asks how that motion "
        "changes the radar observables that a receiver would measure or that a radargram would display."
    )
    doc.add_paragraph("The specific questions are:")
    for item in [
        "How does the parabolic altitude profile change slant range and two-way echo time?",
        "How can an off-track or side-offset target map into an apparent depth in a nadir-style radargram?",
        "What phase signature would a VHF interferometric baseline see for the same geometry?",
        "How does the changing range rate produce Doppler shift at HF and VHF wavelengths?",
        "What PRF is needed to sample the modeled Doppler without simple slow-time aliasing?",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    add_callout(
        doc,
        "What is not being solved",
        "This paper does not prove that a 440 m feature is water, ice, or clutter. It uses the 440 m-scale result "
        "only as an example of how a 25 km side offset at 400 km altitude can appear as a shallow radargram depth."
    )

    doc.add_heading("4. Geometry Model", level=1)
    doc.add_paragraph(
        "The coordinate system is intentionally simple. The variable x is along-track position, y is cross-track "
        "side offset, and z is altitude above the local surface. Closest approach occurs at x = 0, where altitude "
        "is z0. The altitude rises away from closest approach according to a parabola."
    )
    add_equation(doc, "z(x) = z0 + a*x^2")
    add_equation(doc, "a = Delta_z_edge / x_edge^2")
    doc.add_paragraph(
        "Here z0 is the closest-approach altitude, x_edge is the edge of the modeled along-track window, and "
        "Delta_z_edge is the altitude increase at +/-x_edge. For the example case, z0 = 400 km, x_edge = 60 km, "
        "and Delta_z_edge = 4 km, so the platform is 404 km high at x = +/-60 km."
    )
    add_figure(
        doc,
        geometry,
        "Figure 1. Simplified parabolic-motion geometry. Left: altitude is modeled as a local parabola around closest approach. Right: the 25 km cross-track offset is an example side distance, not a claimed measured location.",
    )

    doc.add_heading("5. Derived Radar Observables", level=1)
    doc.add_paragraph(
        "All modeled observables are built from the range between the moving radar and the target. The slant range "
        "includes the along-track separation x, side offset y, and altitude z(x). The nadir-only range is the simpler "
        "case where the target is directly below the spacecraft."
    )
    add_equation(doc, "R_off(x) = sqrt(x^2 + y^2 + z(x)^2)")
    add_equation(doc, "R_nadir(x) = z(x)")
    doc.add_paragraph(
        "The extra path length is R_off - R_nadir. If a radargram assumes the return came from nadir, that extra "
        "path can appear as an apparent depth. Dividing by the ice refractive index n converts the extra path in "
        "free-space range into an approximate in-ice depth scale."
    )
    add_equation(doc, "d_app(x) = (R_off(x) - R_nadir(x)) / n")
    add_equation(doc, "Delta_t(x) = 2*(R_off(x) - R_nadir(x)) / c")
    doc.add_paragraph(
        "The factor of 2 appears in the delay equation because radar delay is two-way travel time. The speed c is "
        "the speed of light in vacuum. The model treats the refractive-index conversion as a simple apparent-depth "
        "mapping, not a full wave-propagation model through variable ice."
    )
    doc.add_paragraph(
        "For interferometry, the cross-track look angle theta is estimated from geometry. The VHF phase depends on "
        "the baseline b, wavelength lambda, and the sine of the look angle. This is why phase is a useful diagnostic: "
        "a side-looking return has a nonzero look angle, while a nadir return is near zero phase in this simplified "
        "geometry."
    )
    add_equation(doc, "theta(x) = atan2(sqrt(x^2 + y^2), z(x))")
    add_equation(doc, "phi(x) = (2*pi/lambda)*b*sin(theta(x))")
    doc.add_paragraph(
        "The simplified power comparison uses only spherical spreading. If the same reflector were observed at two "
        "ranges, the relative two-way range loss scales approximately as the fourth power of range."
    )
    add_equation(doc, "P_rel(x) = (R_nadir(x) / R_off(x))^4")
    add_equation(doc, "P_dB(x) = 10*log10(P_rel(x))")
    doc.add_paragraph(
        "Finally, the Doppler calculation uses the chain rule. Along-track position is related to time by x(t) = v*t. "
        "Because z depends on x, the vertical speed is dz/dt = (dz/dx)*(dx/dt). The range rate then produces a "
        "two-way Doppler shift."
    )
    add_equation(doc, "x(t) = v*t")
    add_equation(doc, "dz/dx = 2*a*x")
    add_equation(doc, "dz/dt = 2*a*x*v")
    add_equation(doc, "dR/dt = (x*dx/dt + z*dz/dt) / R")
    add_equation(doc, "f_D = -2*(dR/dt)/lambda")
    add_equation(doc, "PRF_min ~= 2*max(abs(f_D))")

    doc.add_heading("6. Representative Case", level=1)
    doc.add_paragraph(
        "Table 1 separates the example-model choices from the instrument values motivated by the REASON paper. "
        "This separation is important for explaining the work: some values are chosen to make the model concrete, "
        "while others come from the radar-system context."
    )
    add_table(
        doc,
        ["Quantity", "Value", "Role in the model"],
        [
            ["Closest altitude z0", "400 km", "Example altitude inside the REASON operational altitude range."],
            ["Altitude rise at +/-60 km", "4 km", "Chosen local curvature for the simplified parabolic pass."],
            ["Cross-track offset y", "25 km", "Example side distance; chosen because it gives a 440 m-scale apparent depth at 400 km."],
            ["Platform speed v", "4 km/s", "Representative flyby-speed placeholder for Doppler calculations."],
            ["Ice refractive index n", "1.78", "Converts extra path length into apparent in-ice depth."],
            ["VHF wavelength", "5 m", "REASON-like VHF wavelength used for phase and Doppler."],
            ["HF wavelength", "33.3 m", "REASON-like HF wavelength used for Doppler comparison."],
            ["Interferometry baseline b", "5 m", "Example baseline used in the VHF phase estimate."],
            ["PRF range", "50-3000 Hz", "REASON-like context for whether the modeled Doppler can be sampled."],
        ],
        [2100, 1450, 5810],
    )

    doc.add_heading("7. Results", level=1)
    doc.add_paragraph(
        "At closest approach, the example side-offset return has a slant range slightly longer than the nadir range. "
        "That small geometric difference is enough to create a shallow apparent-depth signature if the return is "
        "displayed on a nadir-style depth axis. The same geometry also produces nonzero interferometric phase."
    )
    add_table(
        doc,
        ["Observable at closest approach", "Value", "Interpretation"],
        [
            ["Extra one-way range", f"{values['extra_range_m']:.1f} m", "Additional slant path relative to a nadir return."],
            ["Apparent depth", f"{values['apparent_depth_m']:.1f} m", "Where the side-offset return would plot on a simplified in-ice depth scale."],
            ["Extra two-way delay", f"{values['extra_delay_us']:.2f} us", "How much later the return arrives compared with nadir."],
            ["Look angle", f"{values['look_angle_deg']:.2f} deg", "Small but nonzero off-nadir angle."],
            ["VHF interferometric phase", f"{values['phase_deg']:.1f} deg", "Phase signature expected from a side-looking return."],
            ["Geometry-only power change", f"{values['power_db']:.3f} dB", "Small spreading-loss difference for this offset and altitude."],
            ["Max modeled VHF Doppler", f"{values['max_vhf_doppler_hz']:.0f} Hz", "Largest short-wavelength Doppler magnitude over the pass."],
            ["Simple VHF PRF floor", f"{values['min_prf_hz']:.0f} Hz", "Approximate PRF needed for simple Nyquist sampling of modeled VHF Doppler."],
        ],
        [3000, 1550, 4810],
    )
    add_figure(
        doc,
        results,
        "Figure 2. Compact results for the representative parabolic pass. Apparent depth, delay, phase, and Doppler all change together because they inherit the same z(x) and R(x) geometry.",
        width_inches=6.05,
    )
    doc.add_paragraph(
        "Figure 2 shows the most important interpretation result: the observables are coupled. The apparent-depth "
        "curve is not an independent curve that must be explained separately from phase or Doppler. They are all "
        "different projections of the same motion and range model. VHF Doppler is larger than HF Doppler because "
        "Doppler magnitude is inversely proportional to wavelength."
    )

    doc.add_heading("8. Interpretation: What Should Happen", level=1)
    doc.add_paragraph(
        "The expected behavior is that radar returns shift smoothly as the object moves through the parabolic path. "
        "Near closest approach, the slant-path difference for the 25 km side-offset example maps to roughly a 440 m "
        "apparent depth. Away from closest approach, both the along-track separation and the parabolic altitude term "
        "change the range, so the plotted return bends rather than staying flat."
    )
    doc.add_paragraph(
        "The phase result should also change systematically. A nadir return should have near-zero cross-track phase "
        "in the simplified model, while a side-offset return should have nonzero phase. This means phase can help "
        "separate a geometry-driven side return from a true nadir subsurface return, although a real interpretation "
        "would also require antenna patterns, clutter simulations, terrain, and noise."
    )
    doc.add_paragraph(
        "The Doppler and PRF result is a timing/sampling result. Doppler does not magically set the PRF by itself, "
        "but it does impose a sampling constraint. In the representative case, VHF reaches a modeled maximum near "
        f"{values['max_vhf_doppler_hz']:.0f} Hz, so a simple slow-time Nyquist argument suggests a PRF above "
        f"{values['min_prf_hz']:.0f} Hz. That value is within the cited 50-3000 Hz context, but actual PRF selection "
        "would also depend on along-track spacing, receive-window timing, range ambiguity, coherent summing, and "
        "data-volume constraints."
    )
    add_callout(
        doc,
        "Plain-language result",
        "If an object moves in a parabola, the radar data should not stay constant. The changing altitude and range "
        "make echoes move in apparent depth, arrive at different times, carry changing phase, and create Doppler "
        "that must be sampled by the PRF."
    )

    doc.add_heading("9. Discussion", level=1)
    doc.add_paragraph(
        "The strength of the simplified model is that it makes the dependencies visible. A professor can follow the "
        "argument from one equation to the next: z(x) controls range, range controls delay and apparent depth, look "
        "angle controls phase, and range rate controls Doppler. This is also why the model is useful for checking "
        "whether a problem viewpoint is right. If an interpretation changes one observable but ignores the others, "
        "it is likely missing part of the geometry."
    )
    doc.add_paragraph(
        "The model also explains why the 25 km and 440 m numbers are linked but not identical kinds of distance. "
        "The 25 km is a horizontal side offset in the example geometry. The approximately 440 m result is not a "
        "physical depth chosen independently; it is the apparent depth that results when the extra slant path from "
        "a 25 km side offset at 400 km altitude is divided by the ice refractive index. In other words, 25 km is "
        "where the example target is to the side; 440 m is where its delayed return would plot on a simplified "
        "depth axis."
    )
    doc.add_paragraph(
        "The REASON paper supports the broader operational motivation because it describes a radar system that must "
        "operate across changing flyby altitude, with adjustable parameters such as PRF, receive-window duration, "
        "and receiver gain. However, the paper does not provide the exact z(t) parabola used here. The parabolic "
        "path is therefore best presented as a local educational or exploratory model."
    )

    doc.add_heading("10. Limitations and Missing Inputs", level=1)
    doc.add_paragraph(
        "Several ingredients are missing before this would become a mission-grade analysis. The simplified model "
        "does not include the actual spacecraft trajectory, attitude, antenna beam pattern, surface roughness, "
        "topography, dielectric layering, scattering strength, thermal noise, onboard processing details, or real "
        "receive-window schedules. It also treats the ice refractive index as one constant value."
    )
    doc.add_paragraph("The most important next inputs are:")
    for item in [
        "A real trajectory and altitude history, preferably from SPICE kernels or an equivalent navigation product.",
        "Spacecraft attitude and antenna pointing so that nadir and off-nadir geometry can be separated correctly.",
        "A surface/topography model to test whether side returns or terrain slopes could generate similar curves.",
        "A dielectric model for the ice shell if apparent depth needs to be converted into physical depth.",
        "Instrument-specific timing and processing settings for each part of the flyby.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("11. Conclusion", level=1)
    doc.add_paragraph(
        "This paper reframes the project around the correct problem: understanding the effects of parabolic motion "
        "on radar sounding observables. The key result is that a single motion model can explain coupled changes "
        "in apparent depth, delay, interferometric phase, Doppler, and PRF requirements. In the representative "
        f"case, a 25 km side offset at 400 km altitude produces an apparent depth of about {values['apparent_depth_m']:.1f} m "
        f"and an extra delay of about {values['extra_delay_us']:.2f} us, while the parabolic path produces VHF Doppler "
        f"magnitudes up to about {values['max_vhf_doppler_hz']:.0f} Hz. The expected outcome is not a single claimed "
        "subsurface interpretation, but a predictable family of geometry-driven radar effects. That is the model's "
        "main value: it gives a clear way to explain what should happen, what each formula means, and what additional "
        "data would be needed before making a stronger scientific claim."
    )

    doc.add_heading("Acknowledgments", level=1)
    doc.add_paragraph(
        "This example draft was prepared from local slide, graph, and formula work in the NASA/REASON radar-analysis "
        "project workspace. It is written as a conference-paper style starting point for classroom explanation and "
        "technical discussion."
    )

    doc.add_heading("References", level=1)
    references = [
        "D. D. Blankenship et al., \"Radar for Europa Assessment and Sounding: Ocean to Near-Surface (REASON),\" Space Science Reviews, 2024, doi: 10.1007/s11214-024-01072-3.",
        "M. I. Skolnik, Radar Handbook, 3rd ed. New York, NY, USA: McGraw-Hill, 2008.",
        "M. A. Richards, Fundamentals of Radar Signal Processing, 2nd ed. New York, NY, USA: McGraw-Hill, 2014.",
    ]
    for ref in references:
        doc.add_paragraph(ref, style="List Number")

    doc.add_heading("Appendix A. Formula Glossary", level=1)
    glossary_rows = [
        ["x", "Along-track position measured relative to closest approach."],
        ["y", "Cross-track side offset; in the example, y = 25 km."],
        ["z(x)", "Altitude as a function of along-track position."],
        ["R_off", "Full slant range to a side-offset target."],
        ["R_nadir", "Range to the point directly below the spacecraft."],
        ["d_app", "Apparent depth caused by extra slant path, not necessarily a real physical layer depth."],
        ["Delta_t", "Extra two-way echo delay relative to nadir."],
        ["theta", "Look angle away from nadir."],
        ["phi", "Interferometric phase for baseline b and wavelength lambda."],
        ["f_D", "Two-way Doppler shift from changing range."],
        ["PRF_min", "Simple sampling floor estimated from twice the maximum Doppler magnitude."],
    ]
    add_table(doc, ["Symbol", "Meaning"], glossary_rows, [1400, 7960])


def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    values = model_values()
    geometry, results = make_figures(values)

    doc = Document()
    configure_document(doc)
    write_paper(doc, values, geometry, results)
    doc.core_properties.title = "A Simplified Parabolic-Motion Model for Interpreting Radar Sounding Observables During a Flyby"
    doc.core_properties.subject = "Conference-style research paper draft"
    doc.core_properties.author = "Student Name (placeholder)"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
